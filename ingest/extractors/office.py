"""Word/Excel extraction — §6.

.docx / .xlsx: native Python libraries (python-docx / openpyxl).
.doc / .xls (legacy): LibreOffice headless, per the §6a validation test.
Known caveat (adopted, not fixed — §6a): a LibreOffice conversion can come
back near-empty with no error and no non-zero exit; we detect that after the
fact via a non-whitespace character-count threshold and flag
`warning_empty_conversion` rather than attempting any remediation.

Password-protected Word/Excel files (OLE-legacy or OOXML) are a hard
exclusion, same policy as encrypted PDFs (§6): detected up front via
msoffcrypto-tool, no decryption attempt, status=encrypted. This was added
after the real corpus run turned up several password-protected legacy .xls
files that LibreOffice's CLI silently fails to convert ("source file could
not be loaded", exit code 0, no output file) — without the up-front check
those surfaced as an opaque `failed`, indistinguishable from real corruption.
"""
from __future__ import annotations

import json
import subprocess
import tempfile
import uuid
from pathlib import Path

import msoffcrypto
import openpyxl
from docx import Document as DocxDocument

from ingest.config import Config
from ingest.extractors import ExtractionResult


def _is_encrypted(path: Path) -> bool:
    try:
        with open(path, "rb") as f:
            return msoffcrypto.OfficeFile(f).is_encrypted()
    except Exception:
        # Not a recognizable OLE/OOXML container, or some other read issue —
        # let the real extractor below surface the concrete error instead.
        return False


def extract_docx(path: Path) -> ExtractionResult:
    if _is_encrypted(path):
        return ExtractionResult(status="encrypted")

    try:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
    except Exception as e:
        return ExtractionResult(status="failed", error_message=f"could not read docx: {e}")

    text = "\n".join(parts)
    return ExtractionResult(status="extracted", text=text)


def extract_xlsx(path: Path) -> ExtractionResult:
    if _is_encrypted(path):
        return ExtractionResult(status="encrypted")

    try:
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    except Exception as e:
        return ExtractionResult(status="failed", error_message=f"could not open xlsx: {e}")

    sheet_names = wb.sheetnames
    blocks = []
    try:
        for name in sheet_names:
            ws = wb[name]
            blocks.append(f"## Sheet: {name}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c for c in cells):
                    blocks.append("\t".join(cells))
    finally:
        wb.close()

    text = "\n".join(blocks)
    return ExtractionResult(status="extracted", text=text, sheet_names=sheet_names)


def _run_soffice(path: Path, target_format: str, config: Config, outdir: Path) -> tuple[bool, str, Path]:
    """Returns (ok, error_message, expected_output_path)."""
    profile_dir = outdir / "profile"
    profile_dir.mkdir(exist_ok=True)

    cmd = [
        config.soffice_binary,
        "--headless",
        "--norestore",
        f"-env:UserInstallation=file://{profile_dir}",
        "--convert-to",
        target_format,
        "--outdir",
        str(outdir),
        str(path),
    ]
    try:
        proc = subprocess.run(
            cmd, capture_output=True, text=True, timeout=config.soffice_timeout_seconds
        )
    except subprocess.TimeoutExpired:
        return False, f"soffice timed out after {config.soffice_timeout_seconds}s", outdir
    except OSError as e:
        return False, f"soffice invocation failed: {e}", outdir

    expected = outdir / (path.stem + "." + target_format)
    if proc.returncode != 0:
        return False, f"soffice exited {proc.returncode}: {proc.stderr.strip()[:500]}", expected
    if not expected.exists():
        detail = proc.stderr.strip()[:500]
        msg = "soffice exited 0 but no output file was created"
        if detail:
            msg += f": {detail}"
        return False, msg, expected

    return True, "", expected


def _check_empty(text: str, config: Config) -> bool:
    """True if the conversion is suspiciously near-empty (§6a caveat)."""
    return len(text.strip()) < config.empty_conversion_min_chars


def extract_doc_legacy(path: Path, config: Config) -> ExtractionResult:
    if _is_encrypted(path):
        return ExtractionResult(status="encrypted")

    with tempfile.TemporaryDirectory(prefix=f"soffice_{uuid.uuid4().hex[:8]}_") as tmpdir:
        outdir = Path(tmpdir)
        ok, err, output_path = _run_soffice(path, "txt", config, outdir)
        if not ok:
            return ExtractionResult(status="failed", error_message=err)

        try:
            text = output_path.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            return ExtractionResult(status="failed", error_message=f"could not read soffice output: {e}")

        if _check_empty(text, config):
            return ExtractionResult(
                status="warning_empty_conversion",
                error_message=(
                    f"LibreOffice .doc->txt conversion produced <{config.empty_conversion_min_chars} "
                    "non-whitespace chars; known caveat, not remediated (§6a)"
                ),
            )

    return ExtractionResult(status="extracted", text=text)


def extract_xls_legacy(path: Path, config: Config) -> ExtractionResult:
    if _is_encrypted(path):
        return ExtractionResult(status="encrypted")

    with tempfile.TemporaryDirectory(prefix=f"soffice_{uuid.uuid4().hex[:8]}_") as tmpdir:
        outdir = Path(tmpdir)
        ok, err, output_path = _run_soffice(path, "csv", config, outdir)
        if not ok:
            return ExtractionResult(status="failed", error_message=err)

        try:
            text = output_path.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            return ExtractionResult(status="failed", error_message=f"could not read soffice output: {e}")

        if _check_empty(text, config):
            return ExtractionResult(
                status="warning_empty_conversion",
                error_message=(
                    f"LibreOffice .xls->csv conversion produced <{config.empty_conversion_min_chars} "
                    "non-whitespace chars; known caveat, not remediated (§6a)"
                ),
            )

    # §6a: LibreOffice's CLI --convert-to csv exports only the first sheet — a known,
    # accepted limitation, not a fallback trigger. sheet_names is left null here because
    # we have no reliable way to enumerate the workbook's sheets via the CLI conversion path.
    # (The pipeline logs XLS_FIRST_SHEET_ONLY as a warning for every successful .xls conversion.)
    return ExtractionResult(status="extracted", text=text)
